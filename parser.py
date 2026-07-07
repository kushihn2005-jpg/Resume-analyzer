"""Extracts raw text from uploaded resume files (PDF or DOCX)."""
import io
from pypdf import PdfReader
import docx


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from a PDF or DOCX resume file."""
    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        return _extract_docx(file_bytes)
    elif lower_name.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {filename}. Use PDF, DOCX, or TXT.")


def _extract_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)
    return "\n".join(text_parts)


def _extract_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also grab text inside tables (common in resumes)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)
